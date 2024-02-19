import csv
import json
from docx import Document
from PyQt6.QtWidgets import (QFileDialog, QMessageBox, QDialog, QVBoxLayout, QPushButton, QHBoxLayout)
from openpyxl import Workbook


# Функция для обработки экспорта данных
def export_data_handler(table):
    # Create a QDialog for user interaction
    export_dialog = QDialog()
    export_dialog.setWindowTitle("Выберите формат экспорта")

    # Set the size of the dialog
    export_dialog.setGeometry(100, 100, 300, 200)
    export_dialog.setStyleSheet(
        "background-image: url('assets/icons/background_navigation_correct.png'); "
        "background-position: center; background-repeat: no-repeat;")
    # Create a layout for the dialog
    dialog_layout = QVBoxLayout(export_dialog)

    # Create the first horizontal layout for the first pair of buttons
    button_layout_1 = QHBoxLayout()

    # Add the first pair of export buttons to the dialog
    button_1 = QPushButton("JSON", export_dialog)
    button_1.setStyleSheet("""
        QPushButton {
            background-color: #f1dbf8;
            border: 2px solid #555;
            border-radius: 15px;
        }
    """)
    button_1.setFixedSize(150, 40)
    button_1.clicked.connect(lambda _, fmt="JSON": handle_export(table, fmt))
    button_layout_1.addWidget(button_1)

    button_2 = QPushButton("CSV", export_dialog)
    button_2.setStyleSheet("""
        QPushButton {
            background-color: #f1dbf8;
            border: 2px solid #555;
            border-radius: 15px;
        }
    """)
    button_2.setFixedSize(150, 40)
    button_2.clicked.connect(lambda _, fmt="CSV": handle_export(table, fmt))
    button_layout_1.addWidget(button_2)

    # Add the first horizontal layout to the main vertical layout
    dialog_layout.addLayout(button_layout_1)

    # Create the second horizontal layout for the second pair of buttons
    button_layout_2 = QHBoxLayout()

    # Add the second pair of export buttons to the dialog
    button_3 = QPushButton("DOCX", export_dialog)
    button_3.setStyleSheet("""
        QPushButton {
            background-color: #f1dbf8;
            border: 2px solid #555;
            border-radius: 15px;
        }
    """)
    button_3.setFixedSize(150, 40)
    button_3.clicked.connect(lambda _, fmt="DOCX": handle_export(table, fmt))
    button_layout_2.addWidget(button_3)

    button_4 = QPushButton("Excel", export_dialog)
    button_4.setStyleSheet("""
        QPushButton {
            background-color: #f1dbf8;
            border: 2px solid #555;
            border-radius: 15px;
        }
    """)
    button_4.setFixedSize(150, 40)
    button_4.clicked.connect(lambda _, fmt="Excel": handle_export(table, fmt))
    button_layout_2.addWidget(button_4)

    # Add the second horizontal layout to the main vertical layout
    dialog_layout.addLayout(button_layout_2)

    # Show the dialog and wait for user interaction
    export_dialog.exec()


def handle_export(table, export_format):
    # Получить данные из таблицы
    model = table.model()
    headers = [model.horizontalHeaderItem(col).text() for col in range(model.columnCount())]
    data = [[model.item(row, col).text() for col in range(model.columnCount())] for row in range(model.rowCount())]

    # Выполнить экспорт в зависимости от выбранного формата
    if export_format == "CSV":
        export_csv(headers, data)
    elif export_format == "JSON":
        export_json(headers, data)
    elif export_format == "DOCX":
        export_docx(headers, data)
    elif export_format == "Excel":
        export_excel(headers, data)
    else:
        QMessageBox.critical(None, "Ошибка", "Выбран недопустимый формат экспорта.")


# Экспорт в формат CSV
def export_csv(headers, data):
    file_dialog = QFileDialog()
    file_path, _ = file_dialog.getSaveFileName(None, "Сохранить CSV файл", "", "CSV файлы (*.csv)")
    if file_path:
        with open(file_path, 'w', newline='') as csv_file:
            csv_writer = csv.writer(csv_file)
            csv_writer.writerow(headers)
            csv_writer.writerows(data)
        QMessageBox.information(None, "Успешный экспорт", f"Данные экспортированы в {file_path}")


# Экспорт в формат JSON
def export_json(headers, data):
    file_dialog = QFileDialog()
    file_path, _ = file_dialog.getSaveFileName(None, "Сохранить JSON файл", "", "JSON файлы (*.json)")
    if file_path:
        export_data = {"headers": headers, "data": data}
        with open(file_path, 'w') as json_file:
            json.dump(export_data, json_file, indent=4)
        QMessageBox.information(None, "Успешный экспорт", f"Данные экспортированы в {file_path}")


# Экспорт в формат DOCX
def export_docx(headers, data):
    file_dialog = QFileDialog()
    file_path, _ = file_dialog.getSaveFileName(None, "Сохранить DOCX файл", "", "DOCX файлы (*.docx)")
    if file_path:
        document = Document()
        document.add_paragraph('Экспорт таблицы', style='Heading1')
        table = document.add_table(rows=1, cols=len(headers))
        table.autofit = True
        # Добавление строки заголовков
        for col, header in enumerate(headers):
            table.cell(0, col).text = header
        # Добавление строк данных
        for row_data in data:
            row_cells = table.add_row().cells
            for col, value in enumerate(row_data):
                row_cells[col].text = value
        document.save(file_path)
        QMessageBox.information(None, "Успешный экспорт", f"Данные экспортированы в {file_path}")


# Экспорт в формат Excel
def export_excel(headers, data):
    file_dialog = QFileDialog()
    file_path, _ = file_dialog.getSaveFileName(None, "Сохранить Excel файл", "", "Excel файлы (*.xlsx)")
    if file_path:
        workbook = Workbook()
        worksheet = workbook.active
        # Добавление строки заголовков
        worksheet.append(headers)
        # Добавление строк данных
        for row_data in data:
            worksheet.append(row_data)
        # Сохранение файла в формате .xlsx
        workbook.save(file_path)
        QMessageBox.information(None, "Успешный экспорт", f"Данные экспортированы в {file_path}")
